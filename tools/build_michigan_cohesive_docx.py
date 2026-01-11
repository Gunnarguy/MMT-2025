#!/usr/bin/env python3
"""Build a cohesive Michigan itinerary DOCX and embed extracted images.

Inputs are kept intentionally simple and local to this repo.
"""

from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document as create_document
from docx.document import Document as DocxDocument
from docx.shared import Inches


def add_bullets(doc: DocxDocument, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(item)
        p.style = "List Bullet"


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument(
        "--out",
        default="output/Girls Trip to Michigan - Cohesive Itinerary.docx",
        help="Output docx path",
    )
    parser.add_argument(
        "--images_dir",
        default="output/michigan_docx_images",
        help="Directory containing extracted images",
    )
    args = parser.parse_args()

    out_path = Path(args.out).expanduser().resolve()
    out_path.parent.mkdir(parents=True, exist_ok=True)

    images_dir = Path(args.images_dir).expanduser().resolve()

    doc = create_document()

    doc.add_heading("Girls Trip to Michigan (Cohesive Itinerary)", level=0)

    doc.add_heading("Trip overview", level=1)
    add_bullets(
        doc,
        [
            "Dates: Converge in Chicago on Aug 28–29; road trip begins Aug 29",
            "Primary loop: Chicago → Grand Rapids → Ludington → Sleeping Bear Dunes → Traverse City → Torch Lake → Charlevoix → Mackinac City / Mackinac Island → Frankenmuth → (Midland/Saginaw/Howell) → (Sarnia/Windsor, Canada) → Detroit → Ann Arbor → Kalamazoo → Home",
        ],
    )

    doc.add_heading("Day-by-day plan", level=1)

    doc.add_heading("Aug 28–29 — Arrivals + Chicago", level=2)
    add_bullets(
        doc,
        [
            "Converge in Chicago (Aug 28/29)",
            "Rental car pickup (PU): Aug 29 — load up and leave by 2:00 PM (earlier preferred)",
            "Rental fee estimate: $570 (Costco/Budget); timing may change depending on arrivals",
        ],
    )

    doc.add_heading("Aug 29 — Chicago → Grand Rapids → Ludington", level=2)
    add_bullets(doc, ["Grand Rapids (approx. 3 hr 18)", "Ludington (approx. 1 hr 45)"])
    doc.add_paragraph(
        "Note: Ludington is the port for the Badger Ferry (car ferry across Lake Michigan to Wisconsin)."
    )
    add_bullets(
        doc,
        [
            "Climb lighthouse",
            "Pine Village",
            "Port Museum",
            "Kayak on the lake",
            "Overnight: Ludington possible (depends on departure time from Chicago)",
        ],
    )

    doc.add_heading("Aug 30 — Ludington → Sleeping Bear Dunes → Traverse City", level=2)
    doc.add_paragraph("Sleeping Bear Dunes (approx. 2 hr)")
    add_bullets(
        doc,
        [
            "Pass required: $25 per car",
            "Site pass: https://www.recreation.gov/sitepass/74294",
        ],
    )
    doc.add_paragraph("For a visit under ~3 hours:")
    add_bullets(
        doc,
        [
            "Drive the 7-mile Pierce Stocking Scenic Drive",
            "Do the Dune Climb (or watch!)",
            "Explore Glen Haven and the Maritime Museum",
            "Dip toes in Loon Lake",
        ],
    )
    doc.add_paragraph("Traverse City (approx. 38 min)")
    add_bullets(
        doc,
        [
            "Wineries / distilleries",
            "City trail",
            "Tour town",
            "Lighthouse",
            "Eat well",
            "Overnight: Traverse City suggested",
        ],
    )

    doc.add_heading(
        "Aug 31 — Traverse City → Torch Lake → Charlevoix → Mackinac City", level=2
    )
    doc.add_paragraph("Torch Lake (approx. 40 min)")
    add_bullets(
        doc,
        [
            "Drive around the lake; consider touring Alden",
            "Clear-bottom kayaks or a boat ride down the lake",
            "Petoskey stone search (see locations below)",
        ],
    )
    doc.add_paragraph("Petoskey stone search locations:")
    add_bullets(
        doc,
        [
            "Bryant Park Beach — 1097 Peninsula Dr, Traverse City, MI 49686",
            "Van’s Beach — 205 Cedar St, Leland, MI 49654",
            "Peterson Park — 10001 E Peterson Park Rd, Northport, MI 49670",
            "Magnus Park — 901 West Lake St, Petoskey, MI 49770",
            "Cross Village Beach — 100 Park Ln, Harbor Springs, MI 49740",
        ],
    )
    doc.add_paragraph("Charlevoix (approx. 23 min)")
    add_bullets(doc, ["Lock", "Castle Farms", "Shop the town", "Lighthouse"])
    doc.add_paragraph("Mackinac City (approx. 1 hr 10)")
    add_bullets(
        doc,
        [
            "Base for Mackinac Island",
            "Overnight: Mackinac City (2 nights suggested)",
        ],
    )

    doc.add_heading("Sep 1–2 — Mackinac Island (day trips)", level=2)
    doc.add_paragraph("Mackinac Island day plan")
    add_bullets(
        doc,
        [
            "Target: 7:00 AM",
            "Park at the ferry dock OR call a shuttle if staying within ~1.5 miles of the ferry dock",
            "If driving to ferry: arrive 30 minutes prior to departure (7:30 AM)",
            "Costs: parking $10; ferry $52 per person round trip",
            "Ferry info: 1-800-638-9892",
            "On-island: ride bikes, tour town, eat well",
            "Return: around 4:00–4:30 PM",
        ],
    )

    doc.add_heading(
        "Sep 3 — Mackinac City → Frankenmuth → (Midland/Saginaw/Howell) → (Sarnia, Canada)",
        level=2,
    )
    doc.add_paragraph("Frankenmuth (approx. 2 hr 50)")
    add_bullets(doc, ["Little Bavaria", "Cute / quaint", "Good food?"])
    doc.add_paragraph("Midland / Saginaw / Howell (approx. 1 hr – 1 hr 30)")
    add_bullets(
        doc, ["Small livable towns", "Howell noted as ‘best overall downtown’!"]
    )
    doc.add_paragraph("Optional: Sarnia, Canada (approx. 2 hr to border)")
    add_bullets(doc, ["Overnight: Sarnia optional"])

    doc.add_heading("Sep 4 — (Sarnia) → Windsor, Canada → Detroit", level=2)
    add_bullets(doc, ["Optional: Windsor, Canada (approx. 2 hr to border)"])
    add_bullets(doc, ["Detroit: meet Julia downtown?", "Overnight: Julia’s"])

    doc.add_heading("Sep 5 — Detroit → Ann Arbor → Kalamazoo → Home", level=2)
    add_bullets(
        doc,
        [
            "Ann Arbor (stop / explore)",
            "Kalamazoo (approx. 1 hr 46): good driving break",
            "Home (approx. 3 hr)",
        ],
    )

    doc.add_heading("Notables / logistics", level=1)

    doc.add_heading("Canada driving (rental car)", level=2)
    add_bullets(
        doc,
        [
            "Vehicles may be driven into Canada with no restrictions.",
            "Notify the rental counter at pickup that you plan to drive into Canada so they can provide a Canadian Non-Resident Insurance Card.",
            "Download: www.avis.ca/content/dam/avis/na/ca/common/pdf-files/abg__canada_non_resident_card.pdf",
            "One additional driver fee may be waived for Costco members at participating locations.",
        ],
    )

    doc.add_heading("Booking", level=2)
    add_bullets(
        doc,
        [
            "Budget Confirmation Number: 38128110US6",
            "Booked pickup: Aug 28",
            "Return: Sep 7",
        ],
    )

    doc.add_heading("Rough cost estimate (per person, divided by 3)", level=2)
    add_bullets(
        doc,
        [
            "Hotels/Airbnbs: $200–250/night, 8 nights → ~$540 each",
            "Meals: $65/day → ~$520 each",
            "Gas: $240 total → ~$80 each",
            "Car rental: $570 total → ~$190 each",
            "Ferry/parking/entry fees: ~$100 total → ~$35 each",
            "Activities: kayaks/boat tours/misc (variable)",
            "Approx total: ~$1300 each (not including airfare)",
        ],
    )

    doc.add_heading("Images from the original DOCX", level=1)
    doc.add_paragraph(
        "These images were extracted from the original ‘Girls Trip to Michigan.docx’ and embedded here so they are explicitly accounted for."
    )

    for idx in [1, 2]:
        img = images_dir / f"img_{idx:02d}.png"
        if img.exists():
            doc.add_paragraph(f"Image {idx}")
            doc.add_picture(str(img), width=Inches(6.5))

    doc.save(str(out_path))
    print(f"Wrote: {out_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
